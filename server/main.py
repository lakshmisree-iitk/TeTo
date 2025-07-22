from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import os
import uuid
import shutil
from datetime import datetime
from typing import List, Optional
import uvicorn
from pydantic import BaseModel
import asyncio
import aiofiles
import aiohttp
from bs4 import BeautifulSoup
import PyPDF2
import io
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv
import json

# Load environment variables
load_dotenv()

app = FastAPI(title="TeTo Learning Platform API")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Google Generative AI
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel('gemini-1.5-flash')

# File storage setup
UPLOADS_DIR = os.path.join(os.path.dirname(__file__), "../uploads")
VIDEOS_DIR = os.path.join(UPLOADS_DIR, "videos")
MATERIALS_DIR = os.path.join(UPLOADS_DIR, "materials")

# Ensure directories exist
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(VIDEOS_DIR, exist_ok=True)
os.makedirs(MATERIALS_DIR, exist_ok=True)

# Store uploaded materials content for AI processing
materials_content = []

# Pydantic models
class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    response: str
    sources: List[dict]
    materials_used: List[str]

# Load existing materials on startup
async def load_existing_materials():
    """Load existing materials from the materials directory"""
    global materials_content
    try:
        for filename in os.listdir(MATERIALS_DIR):
            file_path = os.path.join(MATERIALS_DIR, filename)
            file_ext = os.path.splitext(filename)[1].lower()
            
            content = ""
            if file_ext == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = " ".join([page.extract_text() for page in pdf_reader.pages])
            elif file_ext == '.docx':
                doc = Document(file_path)
                content = " ".join([paragraph.text for paragraph in doc.paragraphs])
            elif file_ext in ['.txt', '.rtf']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            elif file_ext == '.doc':
                content = f"Document file: {filename}"
            
            if content:
                materials_content.append({
                    "filename": filename,
                    "content": content
                })
        
        print(f"Loaded {len(materials_content)} existing materials")
    except Exception as error:
        print(f"Error loading existing materials: {error}")

def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """Extract text content from uploaded file based on file type"""
    try:
        if file_ext == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return " ".join([page.extract_text() for page in pdf_reader.pages])
        elif file_ext == '.docx':
            doc = Document(file_path)
            return " ".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_ext in ['.txt', '.rtf']:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_ext == '.doc':
            return f"Document file: {os.path.basename(file_path)}"
        else:
            return ""
    except Exception as error:
        print(f"Error extracting text from file: {error}")
        return ""

async def search_web(query: str) -> List[dict]:
    """Search the web for relevant information"""
    try:
        search_url = f"https://www.google.com/search?q={query}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        async with aiohttp.ClientSession() as session:
            async with session.get(search_url, headers=headers) as response:
                html = await response.text()
        
        soup = BeautifulSoup(html, 'html.parser')
        results = []
        
        # Extract search results
        for element in soup.select('.g'):
            title_elem = element.find('h3')
            snippet_elem = element.find(class_='VwiC3b')
            link_elem = element.find('a')
            
            if title_elem and snippet_elem and link_elem:
                title = title_elem.get_text()
                snippet = snippet_elem.get_text()
                link = link_elem.get('href')
                
                if title and snippet and link:
                    if link.startswith('/url?q='):
                        link = link[7:].split('&')[0]
                    
                    results.append({
                        "title": title,
                        "snippet": snippet,
                        "link": link
                    })
        
        return results[:5]  # Return top 5 results
    except Exception as error:
        print(f"Web search error: {error}")
        return []

@app.on_event("startup")
async def startup_event():
    """Load existing materials when server starts"""
    await load_existing_materials()

@app.post("/api/upload/video")
async def upload_video(video: UploadFile = File(...)):
    """Upload video file"""
    try:
        if not video.filename:
            raise HTTPException(status_code=400, detail="No video file uploaded")
        
        # Check file extension
        file_ext = os.path.splitext(video.filename)[1].lower()
        allowed_extensions = ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.webm']
        
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail="Only video files are allowed")
        
        # Generate unique filename
        unique_name = f"{int(datetime.now().timestamp() * 1000)}-{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(VIDEOS_DIR, unique_name)
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(video.file, buffer)
        
        video_info = {
            "id": str(uuid.uuid4()),
            "filename": unique_name,
            "originalName": video.filename,
            "path": file_path,
            "size": os.path.getsize(file_path),
            "uploadDate": datetime.now().isoformat(),
            "url": f"/api/videos/{unique_name}"
        }
        
        return {
            "message": "Video uploaded successfully",
            "video": video_info
        }
    except Exception as error:
        print(f"Video upload error: {error}")
        raise HTTPException(status_code=500, detail="Failed to upload video")

@app.post("/api/upload/material")
async def upload_material(material: UploadFile = File(...)):
    """Upload reading material file"""
    try:
        if not material.filename:
            raise HTTPException(status_code=400, detail="No material file uploaded")
        
        # Check file extension
        file_ext = os.path.splitext(material.filename)[1].lower()
        allowed_extensions = ['.pdf', '.doc', '.docx', '.txt', '.rtf']
        
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail="Only PDF, DOC, DOCX, TXT, and RTF files are allowed")
        
        # Generate unique filename
        unique_name = f"{int(datetime.now().timestamp() * 1000)}-{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(MATERIALS_DIR, unique_name)
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(material.file, buffer)
        
        # Extract text content
        content = extract_text_from_file(file_path, file_ext)
        
        material_info = {
            "id": str(uuid.uuid4()),
            "filename": unique_name,
            "originalName": material.filename,
            "path": file_path,
            "size": os.path.getsize(file_path),
            "uploadDate": datetime.now().isoformat(),
            "content": content,
            "url": f"/api/materials/{unique_name}"
        }
        
        # Store content for AI processing
        materials_content.append({
            "filename": material.filename,
            "content": content
        })
        
        return {
            "message": "Material uploaded successfully",
            "material": material_info
        }
    except Exception as error:
        print(f"Material upload error: {error}")
        raise HTTPException(status_code=500, detail="Failed to upload material")

@app.get("/api/videos/{filename}")
async def stream_video(filename: str, request: Request):
    """Stream video file with range support"""
    video_path = os.path.join(VIDEOS_DIR, filename)
    
    if not os.path.exists(video_path):
        raise HTTPException(status_code=404, detail="Video not found")
    
    file_size = os.path.getsize(video_path)
    range_header = request.headers.get("range")
    
    if range_header:
        # Parse range header
        range_str = range_header.replace("bytes=", "")
        start, end = range_str.split("-")
        start = int(start)
        end = int(end) if end else file_size - 1
        chunk_size = end - start + 1
        
        # Return partial content
        headers = {
            "Content-Range": f"bytes {start}-{end}/{file_size}",
            "Accept-Ranges": "bytes",
            "Content-Length": str(chunk_size),
            "Content-Type": "video/mp4",
        }
        
        def video_stream():
            with open(video_path, "rb") as file:
                file.seek(start)
                remaining = chunk_size
                while remaining > 0:
                    chunk_size_read = min(8192, remaining)
                    chunk = file.read(chunk_size_read)
                    if not chunk:
                        break
                    yield chunk
                    remaining -= len(chunk)
        
        return StreamingResponse(
            video_stream(),
            headers=headers,
            status_code=206
        )
    else:
        # Return full file
        headers = {
            "Content-Length": str(file_size),
            "Content-Type": "video/mp4",
        }
        
        def video_stream():
            with open(video_path, "rb") as file:
                while True:
                    chunk = file.read(8192)
                    if not chunk:
                        break
                    yield chunk
        
        return StreamingResponse(
            video_stream(),
            headers=headers
        )

@app.get("/api/materials/{filename}")
async def serve_material(filename: str):
    """Serve material file for download"""
    material_path = os.path.join(MATERIALS_DIR, filename)
    
    if not os.path.exists(material_path):
        raise HTTPException(status_code=404, detail="Material not found")
    
    return FileResponse(material_path)

@app.get("/api/files")
async def get_files():
    """Get list of uploaded files"""
    try:
        videos = []
        for filename in os.listdir(VIDEOS_DIR):
            videos.append({
                "type": "video",
                "filename": filename,
                "url": f"/api/videos/{filename}"
            })
        
        materials = []
        for filename in os.listdir(MATERIALS_DIR):
            materials.append({
                "type": "material",
                "filename": filename,
                "url": f"/api/materials/{filename}"
            })
        
        return {
            "videos": videos,
            "materials": materials
        }
    except Exception as error:
        print(f"Error getting files: {error}")
        raise HTTPException(status_code=500, detail="Failed to get files")

@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """AI Chatbot endpoint"""
    try:
        message = request.message
        
        if not message:
            raise HTTPException(status_code=400, detail="Message is required")
        
        # Search the web for relevant information
        web_results = await search_web(message)
        
        # Prepare context from uploaded materials
        materials_context = ""
        if materials_content:
            materials_context = "\n\nUploaded Materials Context:\n"
            for material in materials_content:
                materials_context += f"{material['filename']}:\n{material['content'][:1000]}...\n\n"
        
        # Prepare web search context
        web_context = ""
        if web_results:
            web_context = "\n\nWeb Search Results:\n"
            for result in web_results:
                web_context += f"{result['title']}:\n{result['snippet']}\nSource: {result['link']}\n\n"
        
        # Create the full context for the AI
        full_context = f"""You are a helpful educational assistant. Answer the student's question based on the following information:

{materials_context}
{web_context}

Student Question: {message}

IMPORTANT INSTRUCTIONS:
1. ALWAYS answer the question FIRST using the uploaded materials if available and relevant
2. If you use information from uploaded materials, cite the filename like this: "According to [filename]: [quote or paraphrase]"
3. After addressing the question with uploaded materials, you can provide additional information from your knowledge
4. If you use web search results, include the exact URL as a hyperlink: "[Title](URL)"
5. Be educational, clear, and comprehensive
6. Always mention which materials you referenced in your response

Answer:"""
        
        # Use Gemini API
        response = model.generate_content(full_context)
        ai_response = response.text
        
        return ChatResponse(
            response=ai_response,
            sources=[{"title": r["title"], "url": r["link"]} for r in web_results],
            materials_used=[m["filename"] for m in materials_content] if materials_content else []
        )
    
    except Exception as error:
        print(f"Chat error: {error}")
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to process chat request: {str(error)}"
        )

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "OK",
        "timestamp": datetime.now().isoformat()
    }

if __name__ == "__main__":
    port = int(os.getenv("PORT", 5001))
    print(f"Server running on port {port}")
    print(f"Upload directories created: {UPLOADS_DIR}")
    uvicorn.run(app, host="0.0.0.0", port=port) 